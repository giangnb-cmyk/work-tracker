"""Doc file tai lieu -> cac (section, text), va cat chunk. Logic thuan (pure), de test.

Ho tro: PDF, Word (.docx), Excel (.xlsx/.xlsm), CSV, TXT, Markdown.
Thu vien nang (pypdf/python-docx/openpyxl) import lazy -> chi can khi thuc su nap dinh dang do.
"""

import csv as _csv
import os

CHUNK_SIZE = 800     # so ky tu moi chunk (xap xi)
CHUNK_OVERLAP = 120  # goi dau giua 2 chunk lien nhau de giu ngu canh


class UnsupportedFormat(Exception):
    """Dinh dang file khong ho tro."""


# --- Cac reader theo dinh dang: tra ve list (section_label, text) ------------

def _read_pdf(path: str) -> list:
    from pypdf import PdfReader

    reader = PdfReader(path)
    out = []
    for i, page in enumerate(reader.pages, start=1):
        txt = (page.extract_text() or "").strip()
        if txt:
            out.append((f"trang {i}", txt))
    return out


def _read_docx(path: str) -> list:
    import docx

    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(parts)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                text += "\n" + " | ".join(cells)
    return [("toan van", text)] if text.strip() else []


def _read_xlsx(path: str) -> list:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        lines = []
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) for c in row if c is not None]
            if vals:
                lines.append(" | ".join(vals))
        if lines:
            out.append((f"sheet '{ws.title}'", "\n".join(lines)))
    wb.close()
    return out


def _read_csv(path: str) -> list:
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            with open(path, newline="", encoding=enc) as f:
                rows = [" | ".join(r) for r in _csv.reader(f) if any(c.strip() for c in r)]
            return [("toan bang", "\n".join(rows))] if rows else []
        except (UnicodeDecodeError, UnicodeError):
            continue
    return []


def _read_text(path: str) -> list:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(path, encoding=enc) as f:
                content = f.read().strip()
            return [("toan van", content)] if content else []
        except (UnicodeDecodeError, UnicodeError):
            continue
    return []


_READERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".xlsx": _read_xlsx,
    ".xlsm": _read_xlsx,
    ".csv": _read_csv,
    ".txt": _read_text,
    ".md": _read_text,
    ".markdown": _read_text,
}

SUPPORTED_EXTS = frozenset(_READERS)


def is_supported(path: str) -> bool:
    return os.path.splitext(path)[1].lower() in _READERS


def read_sections(path: str) -> list:
    """Doc file -> list (section, text). Nem UnsupportedFormat neu khong ho tro."""
    ext = os.path.splitext(path)[1].lower()
    reader = _READERS.get(ext)
    if not reader:
        raise UnsupportedFormat(f"không hỗ trợ định dạng '{ext}'")
    return reader(path)


# --- Cat chunk (thuan) -------------------------------------------------------

def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list:
    """Cat theo doan van, gom toi ~size ky tu, them goi dau. Tra ve list chuoi."""
    text = (text or "").strip()
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks, cur = [], ""
    for p in paragraphs:
        if len(cur) + len(p) + 1 <= size:
            cur = (cur + "\n" + p) if cur else p
            continue
        if cur:
            chunks.append(cur)
        if len(p) > size:  # doan don qua dai -> cat cung
            for i in range(0, len(p), size - overlap):
                chunks.append(p[i:i + size])
            cur = ""
        else:
            cur = p
    if cur:
        chunks.append(cur)

    if overlap <= 0 or len(chunks) <= 1:
        return chunks
    overlapped = [chunks[0]]
    for i in range(1, len(chunks)):
        overlapped.append(chunks[i - 1][-overlap:] + "\n" + chunks[i])
    return overlapped
